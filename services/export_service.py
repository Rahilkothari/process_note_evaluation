from docx import Document
from docx.shared import Inches, Pt
from models.database import ProcessNote
import pandas as pd
import json

class ExportService:
    @staticmethod
    def export_to_docx(note: ProcessNote) -> str:
        doc = Document()
        
        # Title
        title = doc.add_heading('Process Note', 0)
        title.alignment = 1 # Center
        
        # Header Info
        doc.add_heading('Basic Information', level=1)
        doc.add_paragraph(f"Process Name: {note.process_name}")
        doc.add_paragraph(f"Team: {note.team}")
        doc.add_paragraph(f"Version: {note.version}")
        doc.add_paragraph(f"Subject Matter Expert: {note.subject_matter_expert or 'N/A'}")
        doc.add_paragraph(f"Process Owner: {note.process_owner or 'N/A'}")
        
        # Sections
        doc.add_page_break()
        for section in note.sections:
            doc.add_heading(section.section_id, level=1)
            
            if section.content:
                doc.add_paragraph(section.content)
            
            if section.structured_data:
                # Add a table
                data = section.structured_data
                if data and isinstance(data, list) and len(data) > 0:
                    headers = list(data[0].keys())
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = str(header)
                        
                    for row_data in data:
                        row_cells = table.add_row().cells
                        for i, header in enumerate(headers):
                            row_cells[i].text = str(row_data.get(header, ''))
                            
            doc.add_paragraph() # Spacing
            
        filename = f"Process_Note_{note.process_name.replace(' ', '_')}_v{note.version}.docx"
        doc.save(filename)
        return filename
